import docx

doc = docx.Document('files/multipleParagraphs.docx')
print('Paragraph: ', len(doc.paragraphs))

content = []
for i in range(len(doc.paragraphs)):
    print('Content:', doc.paragraphs[i].text)
    # use append and join to get full text
    content.append(doc.paragraphs[i].text)
print('Everything: \n', '\n'.join(content))

# a run object is distinguish with font type (normal, bold, italic)
print('Runs:', len(doc.paragraphs[1].runs))
