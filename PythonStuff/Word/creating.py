import docx

# restyling
doc = docx.Document('files/multipleParagraphs.docx')
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[0].underline = True
doc.save('files/newStyle.docx')

# create a new one
doc = docx.Document()
# adding heading
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
# adding paragraph
doc.add_paragraph('Hello World')
doc.add_picture('files/zophie.png', width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))
doc.save('files/veryNew.docx')
